from io import BytesIO
from pathlib import Path
import re

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch


LOGO_DIR = Path("data/logos")


def clean_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"###\s*", "", text)
    text = re.sub(r"##\s*", "", text)
    text = re.sub(r"#\s*", "", text)
    return text.strip()


def normalize_team_name(team_name: str) -> str:
    return team_name.lower().replace("-", "_").replace(" ", "_")


def find_logo(team_name: str):
    base = normalize_team_name(team_name)
    for ext in [".png", ".jpg", ".jpeg"]:
        path = LOGO_DIR / f"{base}{ext}"
        if path.exists():
            return path
    return None


def export_docx(report_text: str, team_name: str) -> BytesIO:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    logo = find_logo(team_name)

    if logo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(logo), width=Inches(1.0))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{team_name.upper()} SCOUTING REPORT")
    run.bold = True
    run.font.size = Pt(20)

    doc.add_paragraph("")

    for line in report_text.splitlines():
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        plain = clean_markdown(line)

        if plain.isupper() and len(plain) < 60:
            p = doc.add_paragraph()
            r = p.add_run(plain)
            r.bold = True
            r.font.size = Pt(13)

        elif line.startswith("###") or line.startswith("##") or line.startswith("#"):
            p = doc.add_paragraph()
            r = p.add_run(plain)
            r.bold = True
            r.font.size = Pt(13)

        elif line.startswith("- ") or re.match(r"^\d+\.", line):
            doc.add_paragraph(plain, style="List Bullet")

        else:
            p = doc.add_paragraph(plain)
            p.paragraph_format.space_after = Pt(4)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def export_pdf(report_text: str, team_name: str) -> BytesIO:
    output = BytesIO()

    doc = SimpleDocTemplate(
        output,
        pagesize=letter,
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ScoutTitle",
        parent=styles["Title"],
        fontSize=20,
        leading=24,
        spaceAfter=14,
        alignment=1,
    )
    heading_style = ParagraphStyle(
        "ScoutHeading",
        parent=styles["Heading2"],
        fontSize=13,
        leading=16,
        spaceBefore=10,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "ScoutBody",
        parent=styles["BodyText"],
        fontSize=10,
        leading=13,
        spaceAfter=5,
    )

    story = []
    story.append(Paragraph(f"{team_name.upper()} SCOUTING REPORT", title_style))
    story.append(Spacer(1, 0.1 * inch))

    for line in report_text.splitlines():
        line = line.strip()
        if not line:
            story.append(Spacer(1, 0.08 * inch))
            continue

        plain = clean_markdown(line)

        if plain.isupper() and len(plain) < 60:
            story.append(Paragraph(plain, heading_style))
        elif line.startswith("#"):
            story.append(Paragraph(plain, heading_style))
        elif line.startswith("- "):
            story.append(Paragraph("• " + plain[2:], body_style))
        else:
            story.append(Paragraph(plain, body_style))

    doc.build(story)
    output.seek(0)
    return output