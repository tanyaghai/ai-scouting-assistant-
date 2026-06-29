from io import BytesIO
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.pipelines.scouting_context_builder import build_scouting_context
from src.pipelines.player_note_generator import generate_player_note


LOGO_DIR = Path("data/logos")


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for edge in ("top", "left", "bottom", "right"):
        element = tcPr.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tcPr.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_cell_text(cell, text, size=8.5, bold=False, align_center=True):
    cell.text = str(text) if text is not None else ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_border(cell)

    for p in cell.paragraphs:
        if align_center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(size)
            run.bold = bold


def format_team_key(team_name: str) -> str:
    return team_name.lower().replace("-", "_").replace(" ", "_")


def find_logo(team_name: str):
    key = format_team_key(team_name)
    for ext in [".png", ".jpg", ".jpeg"]:
        path = LOGO_DIR / f"{key}{ext}"
        if path.exists():
            return path
    return None


def fmt_num(value, decimals=1):
    if value is None or value == "":
        return ""
    try:
        return round(float(value), decimals)
    except Exception:
        return value


def fmt_pct(value):
    if value is None or value == "":
        return ""
    try:
        return round(float(value) * 100, 1)
    except Exception:
        return value


def get_player_stats(player):
    raw = player.get("raw_stats", {})
    role = player.get("rule_based_profile", {})
    rule_stats = role.get("stats", {})

    merged = {}
    merged.update(rule_stats)
    merged.update(raw)
    return merged


def fg_made_attempted(stats):
    fgm = stats.get("fgm")
    fga = stats.get("fga")

    if fgm is None or fga is None:
        return "—"

    return f"{int(fgm)}-{int(fga)}"


def three_made_attempted(stats):
    made = stats.get("three_made")
    attempts = stats.get("three_attempts")

    if made is None or attempts is None:
        return "—"

    return f"{int(made)}-{int(attempts)}"


def shorten_note(note: str, max_chars: int = 300) -> str:
    note = note.replace("Sentence 1:", "").replace("Sentence 2:", "").strip()
    note = " ".join(note.split())

    if len(note) <= max_chars:
        return note

    return note[:max_chars].rsplit(" ", 1)[0] + "..."


def add_centered_title(doc, team_name, scout_date):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{team_name} Scout")
    run.bold = True
    run.font.size = Pt(30)

    logo = find_logo(team_name)
    if logo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(str(logo), width=Inches(0.9))

    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = d.add_run(scout_date)
    r.bold = True
    r.font.size = Pt(15)


def add_divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 70)
    run.bold = True


def add_simple_table(container, rows):
    table = container.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            set_cell_text(
                table.cell(r_idx, c_idx),
                value,
                size=8.5,
                bold=(r_idx == 0),
            )

    return table


def add_section_title(container, title):
    p = container.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.underline = True
    r.font.size = Pt(11)


def add_bullets(cell, title, bullets):
    p = cell.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.underline = True
    r.font.size = Pt(14)

    for item in bullets:
        p = cell.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        r = p.add_run(f"• {item}")
        r.font.size = Pt(10)


def build_stats_rows(team_stats):
    return [
        ["PPG", fmt_num(team_stats.get("ppg"))],
        ["FG%", fmt_pct(team_stats.get("fg_pct"))],
        ["3PT%", fmt_pct(team_stats.get("three_pct"))],
        ["FT%", fmt_pct(team_stats.get("ft_pct"))],
        ["Rebounding Per Game", fmt_num(team_stats.get("rebounds_per_game"))],
        ["Assists Per Game", fmt_num(team_stats.get("assists_per_game"))],
        ["Turnovers Per Game", fmt_num(team_stats.get("turnovers_per_game"))],
        ["Steals Per Game", fmt_num(team_stats.get("steals_per_game"))],
        ["Blocks Per Game", fmt_num(team_stats.get("blocks_per_game"))],
    ]


def leading_scorers(players, n=3):
    sorted_players = sorted(
        players,
        key=lambda p: get_player_stats(p).get("ppg") or 0,
        reverse=True,
    )

    rows = [["Player", "MPG", "PPG", "FGM-FGA"]]

    for p in sorted_players[:n]:
        stats = get_player_stats(p)
        rows.append([
            p["name"],
            fmt_num(stats.get("mpg")),
            fmt_num(stats.get("ppg")),
            fg_made_attempted(stats),
        ])

    return rows


def top_three_shooters(players, n=3):
    eligible = []

    for p in players:
        role = p.get("rule_based_profile", {})
        stats = get_player_stats(p)

        mpg = stats.get("mpg") or stats.get("minutes_per_game") or 0
        ppg = stats.get("ppg") or 0
        three_pct = stats.get("three_pct") or 0
        three_attempts = stats.get("three_attempts") or 0
        three_rate = stats.get("three_point_rate") or 0
        offensive_profile = role.get("offensive_profile", [])

        is_real_shooter = (
            mpg >= 8
            and ppg >= 3
            and three_attempts >= 25
            and three_pct > 0
            and (
                "Three-point threat" in offensive_profile
                or "3PT-heavy shot profile" in offensive_profile
                or three_rate >= 0.35
            )
        )

        if is_real_shooter:
            eligible.append(p)

    sorted_players = sorted(
        eligible,
        key=lambda p: (
            get_player_stats(p).get("three_pct") or 0,
            get_player_stats(p).get("three_attempts") or 0,
        ),
        reverse=True,
    )

    rows = [["Player", "MPG", "3PM-3PA", "3PT%"]]

    for p in sorted_players[:n]:
        stats = get_player_stats(p)
        rows.append([
            p["name"],
            fmt_num(stats.get("mpg") or stats.get("minutes_per_game")),
            three_made_attempted(stats),
            fmt_pct(stats.get("three_pct")),
        ])

    return rows


def get_keys_to_win(context):
    style = context.get("team_style", {})
    recent = context.get("team_recent_form", {})

    keys = [
        "Fast. Fierce. Fused.",
        "Move the ball. Hit the open man.",
    ]

    keys.extend(style.get("areas_to_attack", [])[:3])
    keys.extend(recent.get("recent_areas_to_attack", [])[:2])

    return keys[:6]


def get_our_defense(context):
    style = context.get("team_style", {})
    recent = context.get("team_recent_form", {})

    defense = []
    defense.extend(style.get("areas_to_attack", [])[:3])
    defense.extend(recent.get("recent_areas_to_attack", [])[:2])

    if not defense:
        defense = [
            "Stay proactive.",
            "Clean up the boards.",
            "Finish possessions.",
        ]

    return defense[:5]


def add_page_one(doc, context, scout_date):
    team_name = context["team_name"]
    players = context["players"]
    team_stats = context.get("team_stats", {})

    add_centered_title(doc, team_name, scout_date)
    add_divider(doc)

    layout = doc.add_table(rows=1, cols=2)
    layout.alignment = WD_TABLE_ALIGNMENT.CENTER

    left = layout.cell(0, 0)
    right = layout.cell(0, 1)

    add_section_title(left, f"{team_name} Statistics")
    add_simple_table(left, build_stats_rows(team_stats))

    left.add_paragraph("")
    add_section_title(left, "Leading Scorers")
    add_simple_table(left, leading_scorers(players))

    left.add_paragraph("")
    add_section_title(left, "Top 3pt Shooters")
    add_simple_table(left, top_three_shooters(players))

    add_bullets(right, "Keys to the Win", get_keys_to_win(context))
    right.add_paragraph("")
    add_bullets(right, "Our Defense", get_our_defense(context))


def add_personnel_pages(doc, context, scout_date):
    doc.add_page_break()

    add_centered_title(doc, context["team_name"], scout_date)
    add_divider(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Personnel")
    r.bold = True
    r.font.size = Pt(14)

    players = context["players"]

    headers = ["Player", "POS.", "MPG", "PPG", "FGM-A", "3PM-3PA", "RPG", "NOTES"]

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, size=8.5, bold=True)

    for player in players[:9]:
        row = table.add_row().cells
        stats = get_player_stats(player)

        try:
            note = shorten_note(generate_player_note(player))
        except Exception:
            note = ""

        values = [
            player["name"],
            "",
            fmt_num(stats.get("mpg") or stats.get("minutes_per_game")),
            fmt_num(stats.get("ppg")),
            fg_made_attempted(stats),
            three_made_attempted(stats),
            fmt_num(stats.get("rpg")),
            note,
        ]

        for i, value in enumerate(values):
            set_cell_text(
                row[i],
                value,
                size=8 if i != 7 else 7.5,
                bold=False,
                align_center=(i != 7),
            )


def export_chapman_style_docx(team_name: str, scout_date: str = None) -> BytesIO:
    if scout_date is None:
        scout_date = date.today().strftime("%m.%d.%y")

    context = build_scouting_context(team_name)

    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    add_page_one(doc, context, scout_date)
    add_personnel_pages(doc, context, scout_date)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


if __name__ == "__main__":
    output = export_chapman_style_docx("Claremont_Mudd_Scripps")

    with open("cms_scout_template.docx", "wb") as f:
        f.write(output.read())

    print("Saved cms_scout_template.docx")